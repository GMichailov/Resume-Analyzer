from utils import _normalize_text

from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
import os
import re
from docx import Document
from pathlib import Path

def read_resume(path: str):
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError("File not found.")
    file_type = p.suffix.lower()
    if file_type == ".pdf":
        return parse_pdf_resume(p)
    elif file_type == ".docx":
        return parse_docx_resume(p)
    else:
        raise ValueError(f"Unsopported file type: Expected pdf or docx.")
    
def parse_pdf_resume(path: Path):
    layout_params = LAParams(line_margin=0.2, word_margin=0.1, char_margin=2.0)
    text = extract_text(str(path), laparams=layout_params) or ""
    return _normalize_text(text)

def parse_docx_resume(path: Path):
    doc = Document(str(path))
    parts=[]

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return _normalize_text("\n".join(parts))